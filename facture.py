from docx import Document

def remplacer_paragraphe(paragraphe, remplacements):
    texte_complet = "".join(run.text for run in paragraphe.runs)
    texte_modifie = texte_complet
    for cle, valeur in remplacements.items():
        texte_modifie = texte_modifie.replace(cle, valeur)

    if texte_modifie != texte_complet:
        for run in paragraphe.runs:
            run.text = ""
        if paragraphe.runs:
            paragraphe.runs[0].text = texte_modifie

def remplacer_dans_cellule(cell, remplacements):
    for paragraphe in cell.paragraphs:
        remplacer_paragraphe(paragraphe, remplacements)

def remplacer_texte(doc, remplacements):
    for paragraphe in doc.paragraphs:
        remplacer_paragraphe(paragraphe, remplacements)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                remplacer_dans_cellule(cell, remplacements)

def main():
    chemin_modele = "modele_facture.docx"
    doc = Document(chemin_modele)

    # Infos client et facture 
    remplacements = {
        "<<NOM_CLIENT>>": input("Nom du client : "),
        "<<ADRESSE_CLIENT>>": input("Adresse du client : "),
        "<<CP_VILLE>>": input("Code postal & Ville : "),
        "<<PAYS>>": input("Pays : "),
        "<<ADRESSE_CHANTIER>>": input("Adresse du chantier : "),
        "<<DT_FACTURE>>": input("Date de la facture : "),
        "<<N_FACTURE>>": input("Numéro de la facture : "),
        "<<ECHEANCE_F>>": input("Date d’échéance de la facture : "),
    }

    total_ht = 0.0 

    print("\nSaisie des lignes de facture d'acompte (jusqu’à 7 lignes) :")
    for i in range(1, 8):
        ajouter = input(f"\nSouhaitez-vous ajouter la ligne {i} ? (oui/non) : ").strip().lower()
        if ajouter == "oui":
            description = input("  Description : ")
            quantite_str = input("  Quantité : ").replace(",", ".")
            unite = input("  Unité : ")
            prix_str = input("  Prix unitaire (€) : ").replace(",", ".")

            try:
                quantite = float(quantite_str)
                prix = float(prix_str)
                montant = round(quantite * prix, 2)
                total_ht += montant
                montant_str = f"{montant:.2f}"
            except ValueError:
                quantite_str = prix_str = montant_str = "Erreur"
        else:
            description = quantite_str = unite = prix_str = montant_str = ""

        remplacements[f"<<DESCRIPTION_{i}>>"] = description
        remplacements[f"<<QTE_{i}>>"] = quantite_str
        remplacements[f"<<U_{i}>>"] = unite
        remplacements[f"<<PU_{i}>>"] = prix_str
        remplacements[f"<<MT_{i}>>"] = montant_str

    remplacements["<<T>>"] = f"{total_ht:.2f}"

    try:
        pourcentage = float(input("\nQuel pourcentage d'acompte ? (ex: 40) : ").replace(",", "."))
        montant_acompte = round(total_ht * (pourcentage / 100), 2)
        remplacements["<<F>>"] = f"{montant_acompte:.2f}"
    except ValueError:
        remplacements["<<F>>"] = "Erreur"

    remplacer_texte(doc, remplacements)

    fichier_sortie = "Factures finales/Facture_Générée.docx"
    doc.save(fichier_sortie)
    print(f"\n Facture générée avec succès : {fichier_sortie}")

if __name__ == "__main__":
    main()
