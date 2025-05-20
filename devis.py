from docx import Document

def remplacer_paragraphe(paragraphe, remplacements):
    texte_complet = "".join(run.text for run in paragraphe.runs)
    texte_modifie = texte_complet
    for cle, valeur in remplacements.items():
        texte_modifie = texte_modifie.replace(cle, valeur)

    if texte_modifie != texte_complet:
        for run in paragraphe.runs:
            run.text = ""
        if paragraphe.runs:
            paragraphe.runs[0].text = texte_modifie

def remplacer_dans_cellule(cell, remplacements):
    for paragraphe in cell.paragraphs:
        remplacer_paragraphe(paragraphe, remplacements)

def remplacer_texte(doc, remplacements):
    for paragraphe in doc.paragraphs:
        remplacer_paragraphe(paragraphe, remplacements)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                remplacer_dans_cellule(cell, remplacements)

def main():
    chemin_modele = "modele_devis.docx"
    doc = Document(chemin_modele)

    remplacements = {
        "<<NOM_CLIENT>>": input("Nom du client : "),
        "<<ADRESSE_CLIENT>>": input("Adresse du client : "),
        "<<CP_VILLE>>": input("Code postal & Ville : "),
        "<<PAYS>>": input("Pays : "),
        "<<ADRESSE_CHANTIER>>": input("Adresse du chantier : "),
        "<<DATE_DEVIS>>": input("Date du devis : "),
        "<<N_DEVIS>>": input("Numéro du devis : "),
        "<<REF>>": input("Référence : "),
    }

    total_ht = 0.0  

    print("\nSaisie des lignes du devis (jusqu’à 10 lignes) :")
    for i in range(1, 11):
        ajouter = input(f"\nSouhaitez-vous ajouter la ligne {i} ? (oui/non) : ").strip().lower()
        if ajouter == "oui":
            description = input("  Description : ")
            quantite_str = input("  Quantité : ").replace(",", ".")
            unite = input("  Unité : ")
            prix_str = input("  Prix unitaire (€) : ").replace(",", ".")

            try:
                quantite = float(quantite_str)
                prix = float(prix_str)
                montant = round(quantite * prix, 2)
                total_ht += montant
                montant_str = f"{montant:.2f}"
            except ValueError:
                quantite_str = prix_str = montant_str = "Erreur"
        else:
            description = quantite_str = unite = prix_str = montant_str = ""

        remplacements[f"<<DESCRIPTION_{i}>>"] = description
        remplacements[f"<<QTE_{i}>>"] = quantite_str
        remplacements[f"<<U_{i}>>"] = unite
        remplacements[f"<<PU_{i}>>"] = prix_str
        remplacements[f"<<MT_{i}>>"] = montant_str

    remplacements["<<T>>"] = f"{total_ht:.2f}"

    remplacer_texte(doc, remplacements)

    fichier_sortie = "Devis finaux/Devis_Généré.docx"
    doc.save(fichier_sortie)
    print(f"\n Devis généré avec succès : {fichier_sortie}")

if __name__ == "__main__":
    main()
